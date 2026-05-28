import sys, io, requests, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from pathlib import Path
from docx import Document

headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36"
}

def get_docx_link(page_url):
    try:
        r = requests.get(page_url, headers=headers, timeout=20)
        if r.status_code != 200:
            return None
        matches = re.findall(
            r'href=["\'](https://thuvienhoclieu\.com/wp-content/uploads/[^"\'<> ]+\.docx)["\']',
            r.text, re.IGNORECASE
        )
        return matches[0] if matches else None
    except Exception as e:
        print(f"Error scraping {page_url}: {e}")
        return None

def check_docx_math(filepath):
    try:
        doc = Document(filepath)
        has_omath = False
        has_object = False
        
        # Check all paragraphs
        for p in doc.paragraphs:
            for child in p._element:
                tag = child.tag.split("}")[-1]
                if tag in ("oMath", "oMathPara"):
                    has_omath = True
                elif tag == "object":
                    has_object = True
                    
        # Check tables
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for child in p._element:
                            tag = child.tag.split("}")[-1]
                            if tag in ("oMath", "oMathPara"):
                                has_omath = True
                            elif tag == "object":
                                has_object = True
                                
        print(f"File: {Path(filepath).name} | Has standard oMath: {has_omath} | Has MathType objects: {has_object}")
    except Exception as e:
        print(f"Error checking {filepath}: {e}")

# Try next Physics and Chemistry candidates
urls = [
    # Physics 2 (Dong Nai)
    'https://thuvienhoclieu.com/de-thi-thu-tot-nghiep-vat-li-2026-so-gd-dong-nai-lan-1-giai-chi-tiet/',
    # Chemistry 1 (Lam Dong)
    'https://thuvienhoclieu.com/de-thi-thu-tot-nghiep-2026-mon-hoa-so-gd-lam-dong-giai-chi-tiet/',
    # Chemistry 2 (Phu Tho)
    'https://thuvienhoclieu.com/de-thi-thu-tot-nghiep-2026-mon-hoa-so-gd-phu-tho-lan-2-giai-chi-tiet/'
]

for url in urls:
    print(f"Scraping link from: {url}")
    link = get_docx_link(url)
    if link:
        filename = link.split("/")[-1]
        dest = Path("data/pdfs") / filename
        if not dest.exists():
            print(f"Downloading {filename}...")
            r = requests.get(link, headers=headers)
            dest.write_bytes(r.content)
            print("Done downloading.")
        check_docx_math(dest)
    else:
        print("No docx link found.")
