"""
processor/azota_parser.py — Parser DOCX cho đề soạn theo định dạng chuẩn Azota.

Đọc deterministic (không LLM) một file .docx giáo viên soạn đúng format Azota:
  - Phần: PHẦN I / PHẦN II / PHẦN III
  - Câu:  "Câu 1." / "Câu 1:" / "Bài 1." / "Question 1."
  - Phần I  (trắc nghiệm):   A. B. C. D.
  - Phần II (đúng/sai):      a) b) c) d)
  - Phần III (trả lời ngắn): "Đáp án:" ngay sau câu (biến thể ngăn bằng "|")
  - Đáp án đúng: bảng "BẢNG ĐÁP ÁN" sau "HẾT" (ưu tiên, để LLM/regex xử lý) HOẶC
                 gạch chân (run.underline) / highlight (run.font.highlight_color)
  - Lời giải:    "Lời giải" / "Hướng dẫn chi tiết" / "Giải thích chi tiết"
  - Mức độ:      tiền tố [0, NB] [1, TH] [2, VD] [3, VDC]

Công thức MathType/WMF được trích thành ảnh PNG inline (xem processor/docx_images.py).
Trả về AzotaExam; việc trích đáp án từ bảng (LLM) + validate cấu trúc nằm ở entry point.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_UNDERLINE
from docx.oxml.ns import qn

from . import docx_images
from .smart_parser import extract_metadata  # tái dùng detect môn/mã đề nếu cần


# ─────────────────────────────────────────────────────────────────────────────
# DATA CLASSES
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class AzotaQuestion:
    index: int                          # số câu trong phần (1-based)
    section: int                        # 1 | 2 | 3
    q_type: str                         # trac_nghiem | dung_sai | tu_luan
    question_text: str                  # đề bài (có thể chứa ![](path))
    options: Optional[dict] = None      # {"A": "...", ...} | {"a": "...", ...}
    correct_answer: Optional[str] = None
    explanation: Optional[str] = None
    level: Optional[str] = None         # Nhận biết | Thông hiểu | Vận dụng | Vận dụng cao
    image_paths: list[str] = field(default_factory=list)
    needs_review: bool = False
    review_reason: str = ""


@dataclass
class AzotaExam:
    subject: str = "UNKNOWN"
    ma_de: str = "01"
    questions: list[AzotaQuestion] = field(default_factory=list)
    raw_answer_block: str = ""          # text vùng đáp án (paragraph + bảng) cho LLM/regex
    fmt_answers: dict = field(default_factory=dict)   # {(section, index): answer} từ format
    diagnostics: dict = field(default_factory=dict)


# ─────────────────────────────────────────────────────────────────────────────
# PATTERNS
# ─────────────────────────────────────────────────────────────────────────────

_RE_SECTION = re.compile(r"^\s*(?:#{0,3}\s*)?PH[ẦÀA]N\s+(I{1,3}|IV|[1-4])\b", re.IGNORECASE)
_RE_QUESTION = re.compile(r"^\s*(?:Câu|Cau|Bài|Bai|Question|Quest)\s*(\d+)\s*[.:]\s*(.*)",
                          re.IGNORECASE | re.DOTALL)
_RE_OPT_ABCD = re.compile(r"^\s*([A-D])[.)]\s")
_RE_SUBITEM = re.compile(r"^\s*([a-d])[.)]\s")
_RE_LEVEL_PREFIX = re.compile(r"^\s*\[\s*([0-3])\s*,\s*(NB|TH|VD|VDC)\s*\]\s*", re.IGNORECASE)
_RE_TLN_ANSWER = re.compile(r"Đáp\s*án\s*[:：]\s*(.+)", re.IGNORECASE)
_RE_MA_DE = re.compile(r"M[ãa]\s*[đd][eề]\s*(?:thi\s*)?[:\s]*?(\d{2,4})", re.IGNORECASE)

# Mốc bắt đầu vùng đáp án / kết thúc đề
_RE_END = re.compile(r"(?:^|[-*\s])-{2,}\s*H[ẾE]T\s*-{2,}|^\s*H[ẾE]T\s*$|B[ẢA]NG\s*Đ[ÁA]P\s*[ÁA]N",
                     re.IGNORECASE | re.MULTILINE)
_RE_SOLUTION = re.compile(
    r"^\s*(?:#{0,3}\s*)?(?:L[ờo]i\s*gi[ảa]i|H[ưươ]ớng\s*d[ẫa]n\s*(?:chi\s*ti[ếe]t|gi[ảa]i)?|"
    r"Gi[ảa]i\s*th[íi]ch\s*chi\s*ti[ếe]t|Đ[áa]p\s*[áa]n\s*[-–]\s*l[ờo]i\s*gi[ảa]i)\b",
    re.IGNORECASE,
)

_LEVEL_CODE = {"NB": "Nhận biết", "TH": "Thông hiểu", "VD": "Vận dụng", "VDC": "Vận dụng cao"}
_SECTION_QTYPE = {1: "trac_nghiem", 2: "dung_sai", 3: "tu_luan"}
_ROMAN = {"I": 1, "II": 2, "III": 3, "IV": 4, "1": 1, "2": 2, "3": 3, "4": 4}


# ─────────────────────────────────────────────────────────────────────────────
# BODY ITERATION (paragraph + table theo đúng thứ tự document)
# ─────────────────────────────────────────────────────────────────────────────

def _iter_body_blocks(doc):
    """Yield ('para', Paragraph) | ('table', Table) theo thứ tự xuất hiện trong body."""
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield ("para", Paragraph(child, doc))
        elif child.tag == qn("w:tbl"):
            yield ("table", Table(child, doc))


def _table_to_text(table) -> str:
    """Serialize 1 bảng Word thành text dạng 'a | b | c' từng dòng (cho LLM/regex đọc)."""
    lines = []
    for row in table.rows:
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        lines.append(" | ".join(cells))
    return "\n".join(lines)


def _looks_like_answer_table(table) -> bool:
    """
    True nếu bảng trông giống BẢNG ĐÁP ÁN (không phải bảng dữ liệu trong câu hỏi).
    Dấu hiệu: header chứa Câu/Đáp án/Mã đề/Chọn, HOẶC là dãy số liên tục, HOẶC sub-item 1a,1b.
    """
    if not table.rows:
        return False
    cells = [c.text.strip() for c in table.rows[0].cells]
    joined = " ".join(cells).lower()
    if re.search(r"c[âa]u|đ[áa]p\s*[áa]n|dap\s*an|m[ãa]\s*đ[eề]|ma\s*de|ch[ọo]n", joined):
        return True
    try:
        from .smart_parser import _is_numeric_header, _is_subitem_header
        if _is_numeric_header(cells[:8]) or _is_subitem_header(cells[:8]):
            return True
    except Exception:
        pass
    return False


# ─────────────────────────────────────────────────────────────────────────────
# FORMAT DETECTION: gạch chân / highlight
# ─────────────────────────────────────────────────────────────────────────────

def _run_is_marked(run) -> bool:
    """True nếu run được gạch chân hoặc highlight (đánh dấu đáp án đúng kiểu Azota)."""
    u = run.underline
    if u is True or (u is not None and u != WD_UNDERLINE.NONE and u is not False):
        return True
    try:
        hl = run.font.highlight_color
        if hl is not None and int(hl) != 0:   # 0 = AUTO/none
            return True
    except Exception:
        pass
    # Shading màu nền (w:shd fill khác auto/trắng)
    rpr = run._element.find(qn("w:rPr"))
    if rpr is not None:
        shd = rpr.find(qn("w:shd"))
        if shd is not None:
            fill = (shd.get(qn("w:fill")) or "").upper()
            if fill and fill not in ("AUTO", "FFFFFF"):
                return True
    return False


def _marked_letters(paragraph, letter_re: re.Pattern) -> set[str]:
    """
    Trả về tập các nhãn (A/B/C/D hoặc a/b/c/d) có nội dung được gạch chân/highlight
    trong paragraph. Theo dõi nhãn hiện tại khi duyệt run theo thứ tự.
    """
    marked: set[str] = set()
    current: Optional[str] = None
    for run in paragraph.runs:
        text = run.text or ""
        # Cập nhật nhãn hiện tại theo các marker xuất hiện trong run này
        for m in re.finditer(r"([A-Da-d])[.)]", text):
            current = m.group(1)
        if current and _run_is_marked(run) and run.text.strip():
            marked.add(current)
    return marked


# ─────────────────────────────────────────────────────────────────────────────
# OPTION / SUBITEM PARSING (từ text đã render kèm ảnh)
# ─────────────────────────────────────────────────────────────────────────────

def _split_options(text: str, upper: bool) -> dict:
    """
    Tách options thành {'A': 'x', ...}. Hỗ trợ marker phân tách bằng đầu dòng,
    newline, tab, hoặc ≥2 khoảng trắng (vd 'A. x\\tB. y' lẫn 'A. x\\nB. y').
    upper=True cho A-D (trắc nghiệm), False cho a-d (đúng/sai).
    """
    cls = "A-D" if upper else "a-d"
    marker = re.compile(r"(?:^|[\n\t]|\s{2,})\s*([%s])[.)]\s" % cls)
    matches = list(marker.finditer(text))
    out: dict[str, str] = {}
    for i, m in enumerate(matches):
        letter = m.group(1)
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        content = re.sub(r"\s+", " ", text[start:end]).strip(" .")
        out[letter] = content
    return out


# ─────────────────────────────────────────────────────────────────────────────
# MAIN PARSE
# ─────────────────────────────────────────────────────────────────────────────

class _QBuilder:
    """Tích lũy trạng thái 1 câu trong lúc duyệt paragraph."""
    def __init__(self, section: int, index: int):
        self.section = section
        self.index = index
        self.stem_parts: list[str] = []
        self.option_text_parts: list[str] = []
        self.option_paras: list = []
        self.tln_answer: Optional[str] = None
        self.explanation_parts: list[str] = []
        self.in_solution = False
        self.level: Optional[str] = None
        self.image_paths: list[str] = []
        self.marked: set[str] = set()

    def finalize(self) -> AzotaQuestion:
        q_type = _SECTION_QTYPE.get(self.section, "trac_nghiem")
        question_text = "\n".join(p for p in self.stem_parts if p.strip()).strip()

        # Phần III: "Đáp án:" có thể nằm inline cùng dòng câu hỏi → trích + bỏ khỏi đề
        if self.section == 3:
            if self.tln_answer is None:
                mt = _RE_TLN_ANSWER.search(question_text)
                if mt:
                    self.tln_answer = _normalize_tln(mt.group(1))
            question_text = _RE_TLN_ANSWER.sub("", question_text).strip(" .\n")

        opt_text = "\n".join(self.option_text_parts).strip()
        options: Optional[dict] = None
        if self.section == 2:
            options = _split_options(opt_text, upper=False) or None
        elif self.section != 3:
            options = _split_options(opt_text, upper=True) or None
        # Phần III: không có options

        explanation = "\n".join(p for p in self.explanation_parts if p.strip()).strip(" .\n") or None

        return AzotaQuestion(
            index=self.index,
            section=self.section,
            q_type=q_type,
            question_text=question_text,
            options=options,
            correct_answer=self.tln_answer,   # P3: từ "Đáp án:"; P1/P2 điền sau từ bảng
            explanation=explanation,
            level=self.level,
            image_paths=list(self.image_paths),
        )


def parse_azota_docx(
    path: str,
    image_dir: Optional[str] = None,
    image_prefix: str = "azota",
) -> AzotaExam:
    """Parse 1 file .docx soạn theo format Azota → AzotaExam."""
    doc = Document(path)
    if image_dir is None:
        image_dir = str(Path(path).parent / "azota_images")
    Path(image_dir).mkdir(parents=True, exist_ok=True)

    exam = AzotaExam()
    img_counter = [0]

    # ── Mã đề + môn từ header (≈ 40 paragraph đầu) ──────────────────────────
    header_text = "\n".join(p.text for p in doc.paragraphs[:40])
    m = _RE_MA_DE.search(header_text)
    if m:
        exam.ma_de = m.group(1)
    try:
        meta = extract_metadata(header_text)
        exam.subject = meta.get("subject", "UNKNOWN")
        if meta.get("ma_de") and not m:
            exam.ma_de = meta["ma_de"]
    except Exception:
        pass

    section = 0
    builder: Optional[_QBuilder] = None
    in_answer_region = False
    answer_parts: list[str] = []
    seen_question = False

    def make_saver(b: _QBuilder):
        def _saver(data: bytes, ext: str) -> Optional[str]:
            img_counter[0] += 1
            name = f"{image_prefix}_{img_counter[0]}.{ext}"
            abspath = str(Path(image_dir) / name)
            try:
                with open(abspath, "wb") as f:
                    f.write(data)
            except Exception:
                return None
            b.image_paths.append(abspath)
            return f"![]({abspath})"
        return _saver

    def flush():
        nonlocal builder
        if builder is not None:
            q = builder.finalize()
            # đáp án từ format (gạch chân/highlight)
            if builder.marked:
                if builder.section == 2:
                    code = "".join("D" if c in builder.marked else "S" for c in "abcd")
                    exam.fmt_answers[(2, builder.index)] = code
                else:
                    letters = sorted(L for L in builder.marked if L in "ABCD")
                    if letters:
                        exam.fmt_answers[(builder.section, builder.index)] = "".join(letters)
            exam.questions.append(q)
            builder = None

    for kind, obj in _iter_body_blocks(doc):
        if kind == "table":
            if in_answer_region or (seen_question and _looks_like_answer_table(obj)):
                # Bảng đáp án thật (có header Câu/Đáp án/Mã đề/số liên tục) → vào vùng đáp án
                answer_parts.append(_table_to_text(obj))
                in_answer_region = True
                flush()
            elif builder is not None:
                # Bảng dữ liệu nằm TRONG câu hỏi → render vào đề bài
                builder.stem_parts.append(_table_to_text(obj))
            continue

        para = obj
        text = para.text.strip()
        if not text and not docx_images.has_inline_image(para):
            continue

        # ── Vùng đáp án (sau HẾT / BẢNG ĐÁP ÁN) ────────────────────────────
        if not in_answer_region and _RE_END.search(text):
            flush()
            in_answer_region = True
            # phần sau marker trên cùng dòng (nếu có) cũng đưa vào answer
            answer_parts.append(text)
            continue
        if in_answer_region:
            answer_parts.append(_render(para, _noop_saver))
            continue

        # ── Section header ─────────────────────────────────────────────────
        ms = _RE_SECTION.match(text)
        if ms:
            flush()
            section = _ROMAN.get(ms.group(1).upper(), section + 1)
            continue

        # ── Câu mới ────────────────────────────────────────────────────────
        mq = _RE_QUESTION.match(text)
        if mq:
            flush()
            seen_question = True
            if section == 0:
                section = 1  # đề không ghi PHẦN (vd Tiếng Anh) → mặc định phần 1
            idx = int(mq.group(1))
            builder = _QBuilder(section, idx)
            rest = mq.group(2).strip()
            # tiền tố mức độ [x, XX]
            ml = _RE_LEVEL_PREFIX.match(rest)
            if ml:
                builder.level = _LEVEL_CODE.get(ml.group(2).upper())
                rest = rest[ml.end():].strip()
            # render lại cả paragraph để giữ ảnh trong stem, bỏ prefix "Câu N."
            rendered = _render(para, make_saver(builder))
            rendered = _strip_question_prefix(rendered)
            if ml:
                rendered = _RE_LEVEL_PREFIX.sub("", rendered, count=1)
            builder.stem_parts.append(rendered)
            continue

        if builder is None:
            continue

        # ── Trong 1 câu: lời giải / option / sub-item / đáp án ngắn / nội dung ─
        if _RE_SOLUTION.match(text):
            builder.in_solution = True
            after = _RE_SOLUTION.sub("", text, count=1).strip(" .:–-")
            if after:
                builder.explanation_parts.append(after)
            continue

        if builder.in_solution:
            builder.explanation_parts.append(_render(para, make_saver(builder)))
            continue

        # Phần III: bắt "Đáp án: ..."
        if builder.section == 3:
            mt = _RE_TLN_ANSWER.search(text)
            if mt:
                builder.tln_answer = _normalize_tln(mt.group(1))
                continue

        # Option A./B./C./D. (P1) hoặc sub-item a)/b)/c)/d) (P2)
        is_opt = (builder.section != 3) and (
            _RE_OPT_ABCD.match(text) or _RE_SUBITEM.match(text)
            or bool(re.search(r"\t\s*[A-Da-d][.)]", text))
        )
        if is_opt:
            rendered = _render(para, make_saver(builder))
            builder.option_text_parts.append(rendered)
            builder.option_paras.append(para)
            letter_re = _RE_SUBITEM if builder.section == 2 else _RE_OPT_ABCD
            builder.marked |= _marked_letters(para, letter_re)
            continue

        # Còn lại: nối vào đề bài
        builder.stem_parts.append(_render(para, make_saver(builder)))

    flush()
    exam.raw_answer_block = "\n".join(p for p in answer_parts if p.strip()).strip()

    exam.diagnostics = {
        "n_questions": len(exam.questions),
        "n_sections": len({q.section for q in exam.questions}),
        "has_answer_block": bool(exam.raw_answer_block),
        "n_fmt_answers": len(exam.fmt_answers),
    }
    return exam


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _render(paragraph, saver) -> str:
    return docx_images.render_paragraph(paragraph, saver)


def _noop_saver(data: bytes, ext: str) -> Optional[str]:
    return None  # vùng đáp án: bỏ ảnh, chỉ lấy text


def _strip_question_prefix(text: str) -> str:
    return re.sub(r"^\s*(?:Câu|Cau|Bài|Bai|Question|Quest)\s*\d+\s*[.:]\s*", "", text,
                  count=1, flags=re.IGNORECASE)


def _normalize_tln(raw: str) -> str:
    """'1,5 | 1.5' → '1,5' (lấy biến thể đầu, bỏ khoảng trắng thừa)."""
    first = raw.split("|")[0]
    return re.sub(r"\s+", " ", first).strip(" .")


# ─────────────────────────────────────────────────────────────────────────────
# FORMAT CHECK
# ─────────────────────────────────────────────────────────────────────────────

def is_azota_format(path: str) -> tuple[bool, dict]:
    """
    Kiểm tra file .docx có đúng format Azota không (đủ để route pipeline).

    Tiêu chí: có ≥ 5 câu "Câu N." VÀ (có header PHẦN HOẶC có nguồn đáp án: bảng/HẾT/gạch chân).
    Trả (ok, diagnostics).
    """
    try:
        doc = Document(path)
    except Exception as e:
        return False, {"error": str(e)}

    n_questions = 0
    n_sections = 0
    has_end = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if _RE_QUESTION.match(t):
            n_questions += 1
        if _RE_SECTION.match(t):
            n_sections += 1
        if _RE_END.search(t):
            has_end = True

    has_answer_table = len(doc.tables) > 0
    diag = {
        "n_questions": n_questions,
        "n_sections": n_sections,
        "has_end_marker": has_end,
        "has_table": has_answer_table,
    }
    ok = n_questions >= 5 and (n_sections >= 1 or has_end or has_answer_table)
    diag["is_azota"] = ok
    return ok, diag
