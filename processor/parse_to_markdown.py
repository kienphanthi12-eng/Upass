import sys
import os
import re

# Set encoding to UTF-8
sys.stdout.reconfigure(encoding='utf-8')

def parse_docx(file_path):
    import docx
    doc = docx.Document(file_path)
    md_lines = []

    # Parse paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        
        # Check if there is bold text at the beginning (e.g. "Câu 1:")
        if p.runs and p.runs[0].bold:
            # Format nicely as bold
            text = f"**{text}**"
        
        md_lines.append(text)

    # Convert tables to markdown / HTML tables
    if doc.tables:
        md_lines.append("\n\n## ĐÁP ÁN VÀ BẢNG THÔNG TIN\n")
        for table in doc.tables:
            html_table = ["<table>"]
            for row in table.rows:
                html_table.append("  <tr>")
                for cell in row.cells:
                    # Clean cell text
                    cell_text = cell.text.strip().replace("\n", "<br>")
                    html_table.append(f"    <td>{cell_text}</td>")
                html_table.append("  </tr>")
            html_table.append("</table>\n")
            md_lines.append("\n".join(html_table))

    return "\n\n".join(md_lines)

def parse_latex(file_path):
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()

    # Strip latex boilerplate like preamble, document environment
    # but keep the main body
    body_match = re.search(r'\\begin\{document\}([\s\S]*?)\\end\{document\}', content)
    if body_match:
        body = body_match.group(1)
    else:
        body = content

    # Clean up common LaTeX commands to make it look like Markdown
    # Keep math equations ($...$ and $$...$$) intact as they are compatible with KaTeX/Markdown
    body = re.sub(r'\\maketitle', '', body)
    body = re.sub(r'\\title\{([^\}]+)\}', r'# \1\n', body)
    body = re.sub(r'\\section\{([^\}]+)\}', r'## \1\n', body)
    body = re.sub(r'\\subsection\{([^\}]+)\}', r'### \1\n', body)
    
    # Environment conversions
    body = re.sub(r'\\begin\{question\}', '\n', body)
    body = re.sub(r'\\end\{question\}', '\n', body)
    body = re.sub(r'\\begin\{enumerate\}', '', body)
    body = re.sub(r'\\end\{enumerate\}', '', body)
    body = re.sub(r'\\item', '\n-', body)
    
    # Clean LaTeX symbols
    body = body.replace(r'\\', '\n')
    body = body.replace(r'\_', '_')
    body = body.replace(r'\%', '%')
    body = body.replace(r'\&', '&')
    
    # Normalize multiple newlines
    body = re.sub(r'\n{3,}', '\n\n', body)
    return body.strip()

def main():
    if len(sys.argv) < 2:
        print("Usage: python parse_to_markdown.py <file_path>")
        sys.exit(1)

    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(f"Error: File {file_path} not found")
        sys.exit(1)

    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == ".docx":
            markdown = parse_docx(file_path)
        elif ext in (".tex", ".latex"):
            markdown = parse_latex(file_path)
        else:
            print(f"Error: Unsupported file format {ext}")
            sys.exit(1)
        
        # Output result to stdout
        print(markdown)
    except Exception as e:
        print(f"Error parsing file: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
