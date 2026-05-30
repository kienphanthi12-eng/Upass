"""
import_exam_docx.py — Pipeline DOCX trực tiếp (không cần MinerU/OCR)

Dùng cho: Vật Lý, Hóa Học, Lịch Sử (18 TN + 4 Đ/S + 6 TLN = 28 câu)
          Tiếng Anh (40 TN thuần)

Nguồn DOCX: thuvienhoclieu.com (wp-content/uploads/*.docx)

Usage:
  from import_exam_docx import run_docx_pipeline
  exam_id = run_docx_pipeline(docx_path, title, year, subject_id)
"""
import json
import re
import sys
import io
from pathlib import Path
from typing import Optional, Dict, List, Tuple

if sys.platform == "win32" and hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

from docx import Document
from rich.console import Console
from rich.panel import Panel

sys.path.insert(0, str(Path(__file__).parent))
from database import db

console = Console()

# Offset question_number theo loại
TN_OFFSET  = 0    # 1–40
DS_OFFSET  = 100  # 101–104
TLN_OFFSET = 200  # 201–206


# ─── Answer extraction ────────────────────────────────────────────────────────

def _extract_answers(tables) -> Dict[str, Dict[str, str]]:
    """Trích đáp án từ các bảng đáp án trong DOCX.

    Returns:
        {
          'I':   {'1': 'A', '2': 'B', ...},
          'II':  {'1': 'DSDS', '2': 'DDSS', ...},
          'III': {'1': 'text', ...},
        }
    """
    ans: Dict[str, Dict[str, str]] = {'I': {}, 'II': {}, 'III': {}}

    for table in tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows or len(rows) < 2:
            continue

        header = [c.lower() for c in rows[0]]

        # ── Phần II (Đúng/Sai): header = [câu, a, b, c, d]
        if len(header) >= 5 and header[0] in ('câu', '') and set(header[1:5]) <= {'a', 'b', 'c', 'd', ''}:
            # Check second row has 'câu N' format
            second = rows[1][0].lower()
            if re.search(r'\d', second) or 'câu' in second:
                for row in rows[1:]:
                    m = re.search(r'\d+', row[0])
                    if not m:
                        continue
                    n = m.group()
                    vals = row[1:5]
                    code = ''
                    for v in vals:
                        v = v.strip()
                        if v.upper() in ('Đ', 'D', 'ĐÚNG'):
                            code += 'D'
                        elif v.upper() in ('S', 'SAI'):
                            code += 'S'
                        else:
                            code += '?'
                    ans['II'][n] = code
                continue

        # ── Format (Câu, Đáp án) × n cột — pairs side by side
        # header = ['câu', 'đáp án', 'câu', 'đáp án', ...] (2, 4, 6, 8 cột)
        if (len(header) >= 2 and len(header) % 2 == 0
                and header[0] in ('câu', '')
                and header[1] in ('đáp án', 'answer', 'đáp án (answer)', '')):
            n_pairs = len(header) // 2
            valid = all(
                header[j*2] in ('câu', '') and header[j*2+1] in ('đáp án', 'answer', 'đáp án (answer)', '')
                for j in range(n_pairs)
            )
            if valid:
                for row in rows[1:]:
                    for p in range(n_pairs):
                        num_cell = row[p*2].strip() if p*2 < len(row) else ''
                        ans_cell = row[p*2+1].strip() if p*2+1 < len(row) else ''
                        if not num_cell or not ans_cell:
                            continue
                        m = re.search(r'\d+', num_cell)
                        if not m:
                            continue
                        n = m.group()
                        av = ans_cell.upper()
                        if av in ('A', 'B', 'C', 'D'):
                            ans['I'][n] = av
                        elif re.match(r'^[DS?]{4}$', av):
                            ans['II'][n] = av
                        else:
                            # Text answer (TLN) → Phần III
                            ans['III'][n] = ans_cell
                continue

        # ── Format "N. X" — cells chứa số và đáp án, không có header
        # Ví dụ: ['1. C', '2. C', '3. B', ...] hoặc ['1. C', '2. D', ...]
        if rows and all(
            re.match(r'^\d+\.\s*[A-D]$', c.strip()) or c.strip() == ''
            for row in rows for c in row
            if c.strip()
        ):
            for row in rows:
                for cell in row:
                    m = re.match(r'^(\d+)\.\s*([A-D])$', cell.strip())
                    if m:
                        ans['I'][m.group(1)] = m.group(2)
            continue

        # ── Format plain 2-row: row0 = số câu, row1 = đáp án (không có header)
        # Ví dụ (LY Ca Mau Phần I): row0=['1','2','3',...,'18'], row1=['B','D','A',...,'C']
        # Ví dụ (LY Ca Mau Phần II): row0=['1a','1b','1c','1d','2a',...], row1=['S','Đ','Đ',...]
        if len(rows) == 2 and len(rows[0]) >= 3:
            r0 = rows[0]; r1 = rows[1]
            # Case A: all cells in row0 are plain numbers
            if all(re.match(r'^\d+$', c.strip()) for c in r0 if c.strip()):
                # Case A1: row1 are A/B/C/D → Phần I TN answers
                if all(c.strip().upper() in ('A','B','C','D','') for c in r1):
                    for num, val in zip(r0, r1):
                        n = num.strip(); v = val.strip().upper()
                        if n and v in ('A','B','C','D'):
                            ans['I'][n] = v
                    continue
                # Case A2: row1 are text (numeric or mixed) → Phần III TLN answers
                if all(c.strip() for c in r1):  # all non-empty
                    for num, val in zip(r0, r1):
                        n = num.strip(); v = val.strip()
                        if n and v:
                            ans['III'][n] = v
                    continue
            # Case B: row0 = ['Na','Nb','Nc','Nd','(N+1)a',...], row1 = ['S'/'Đ',...]
            # Composite DS format — rebuild per-question DSDS string
            if all(re.match(r'^\d+[abcd]$', c.strip(), re.IGNORECASE) for c in r0 if c.strip()):
                ds_map: Dict[str, list] = {}
                for cell, val in zip(r0, r1):
                    m = re.match(r'^(\d+)([abcd])$', cell.strip(), re.IGNORECASE)
                    if not m:
                        continue
                    n, letter = m.group(1), m.group(2).lower()
                    if n not in ds_map:
                        ds_map[n] = {'a': '?', 'b': '?', 'c': '?', 'd': '?'}
                    v = val.strip().upper()
                    if v in ('Đ', 'D', 'ĐÚNG', 'TRUE'):
                        ds_map[n][letter] = 'D'
                    elif v in ('S', 'SAI', 'FALSE'):
                        ds_map[n][letter] = 'S'
                for n, parts in ds_map.items():
                    code = parts['a'] + parts['b'] + parts['c'] + parts['d']
                    ans['II'][n] = code
                continue

        # ── Phần I / III: header = [câu / đáp án, 1, 2, 3, ...]
        # Table có nhiều hơn 1 row với pattern câu/đáp án
        i = 0
        while i < len(rows):
            row = rows[i]
            first = row[0].lower() if row else ''
            if first in ('câu', 'câu hỏi', 'câu (question)') and len(row) > 2:
                # header row, next row = answers
                if i + 1 < len(rows):
                    nums = row[1:]
                    vals = rows[i + 1][1:]
                    _ans_label = rows[i + 1][0].lower()
                    # Nếu row[i+1][0] ~ 'đáp án' hoặc '' → đây là hàng giá trị
                    for n, v in zip(nums, vals):
                        n = n.strip()
                        v = v.strip()
                        if not n or not v:
                            continue
                        if v in ('A', 'B', 'C', 'D'):
                            ans['I'][n] = v
                        elif re.match(r'^[DS?]{4}$', v.upper()):
                            ans['II'][n] = v.upper()
                        else:
                            # Text answer → Phần III
                            ans['III'][n] = v
                    i += 2
                    continue
            i += 1

    return ans


# ─── Options parsing ──────────────────────────────────────────────────────────

def _parse_options_text(text: str) -> Dict[str, str]:
    """Phân tích đáp án A/B/C/D từ một đoạn văn bản.

    Hỗ trợ:
    - Tab-separated: 'A. opt\tB. opt\tC. opt\tD. opt'
    - Space-separated: 'A. opt  B. opt  C. opt  D. opt'
    - Single option: 'A. opt'
    """
    opts = {}
    # Tách theo tab hoặc bắt đầu mới bằng [A-D].
    parts = re.split(r'\t+|\s{3,}(?=[A-D][\.\)])', text)
    for part in parts:
        m = re.match(r'^([A-D])[\.\)]\s*(.*)', part.strip(), re.DOTALL)
        if m:
            key = m.group(1)
            val = m.group(2).strip()
            opts[key] = val
    return opts


# ─── Question parsing ─────────────────────────────────────────────────────────

def _save_current(questions: list, q: Optional[dict], opts: dict, stmts: dict):
    """Lưu câu hỏi hiện tại vào danh sách (nếu hợp lệ)."""
    if not q:
        return
    section = q.get('section')
    if not section:
        return

    q['opts'] = dict(opts)
    q['stmts'] = dict(stmts)
    questions.append(q)


def parse_docx_questions(docx_path: str) -> Tuple[str, Dict, List[dict]]:
    """Parse DOCX exam file.

    Returns:
        (exam_title, answers_dict, questions_list)

    questions_list item:
        {
          'section': 'I'|'II'|'III',
          'number': int,   # 1-based within section
          'content': str,
          'opts': dict,    # {'A': ..., 'B': ..., 'C': ..., 'D': ...} for TN/DS
          'stmts': dict,   # {'a': ..., 'b': ..., 'c': ..., 'd': ...} for DS
        }
    """
    doc = Document(docx_path)

    # Đáp án từ bảng
    answers = _extract_answers(doc.tables)

    # Tiêu đề từ bảng đầu tiên
    title = ''
    for table in doc.tables[:2]:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if any(kw in t.upper() for kw in ('KỲ THI', 'ĐỀ THI', 'MÔN THI', 'SỞ GIÁO DỤC')):
                    title = re.sub(r'\s+', ' ', t)[:300]
                    break
            if title:
                break
        if title:
            break

    if not title and doc.paragraphs:
        title = doc.paragraphs[0].text.strip()[:200]

    # Parse paragraphs
    questions: List[dict] = []
    section: Optional[str] = None
    current_q: Optional[dict] = None
    opt_buf: Dict[str, str] = {}
    stmt_buf: Dict[str, str] = {}

    # Patterns for "answer key / end of exam" detection.
    # Use startswith for word-based keywords to avoid false matches like
    # "đáp án của mình" inside exam instructions.
    STOP_STARTS = ('ĐÁP ÁN', 'LỜI GIẢI')
    STOP_CONTAINS = ('----- HẾT', '--- HẾT', 'HẾT -----', '----------\nHẾT')

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        upper = text.upper()

        # Stop at answer key section
        if (any(upper.startswith(kw) for kw in STOP_STARTS)
                or any(kw in upper for kw in STOP_CONTAINS)):
            _save_current(questions, current_q, opt_buf, stmt_buf)
            break

        # Detect section header
        if re.match(r'^PHẦN\s+I[\s\.\:]', text) and 'PHẦN II' not in text:
            _save_current(questions, current_q, opt_buf, stmt_buf)
            current_q = None; opt_buf = {}; stmt_buf = {}
            section = 'I'
            continue
        if re.match(r'^PHẦN\s+II[\s\.\:]', text) and 'PHẦN III' not in text:
            _save_current(questions, current_q, opt_buf, stmt_buf)
            current_q = None; opt_buf = {}; stmt_buf = {}
            section = 'II'
            continue
        if re.match(r'^PHẦN\s+III[\s\.\:]', text):
            _save_current(questions, current_q, opt_buf, stmt_buf)
            current_q = None; opt_buf = {}; stmt_buf = {}
            section = 'III'
            continue

        # New question — Câu N. (Vietnamese) or Question N. (English)
        # Check này trước `if not section` để tiếng Anh tự set section='I'
        m_cau = re.match(r'^Câu\s+(\d+)\s*[\.\:]\s*(.*)', text, re.DOTALL)
        m_qen = re.match(r'^Question\s+(\d+)[\.\:]\s*(.*)', text, re.DOTALL | re.IGNORECASE) if not m_cau else None
        if m_cau or m_qen:
            _save_current(questions, current_q, opt_buf, stmt_buf)
            m = m_cau or m_qen
            qnum = int(m.group(1))
            rest = m.group(2).strip()
            # For English inline format: "Question N. A. opt  B. opt  C. opt  D. opt"
            inline_opts = _parse_options_text(rest) if rest else {}
            if not section:
                section = 'I'  # English exams: auto-set section
            # content = rest (options text as fallback) keeps len > 15 → không bị auto_hide
            current_q = {
                'section': section,
                'number': qnum,
                'content': rest if rest else text,
            }
            opt_buf = inline_opts
            stmt_buf = {}
            continue

        if not section:
            continue

        if not current_q:
            continue

        # Section-specific processing
        if section == 'I':
            # Check if this is an options paragraph
            if re.match(r'^[A-D][\.\)]', text):
                opts = _parse_options_text(text)
                opt_buf.update(opts)
            elif '\t' in text and re.search(r'[A-D][\.\)]', text):
                opts = _parse_options_text(text)
                if opts:
                    opt_buf.update(opts)
                else:
                    current_q['content'] += '\n' + text
            else:
                current_q['content'] += '\n' + text

        elif section == 'II':
            # Check if this is a statement (a/b/c/d)
            m_stmt = re.match(r'^([abcd])[\)\.\s]\s*(.*)', text, re.DOTALL)
            if m_stmt:
                stmt_buf[m_stmt.group(1)] = m_stmt.group(2).strip()
            else:
                current_q['content'] += '\n' + text

        elif section == 'III':
            # Shared context (e.g. "Dùng thông tin sau cho câu N và câu M")
            m_context = re.match(r'^Dùng thông tin', text, re.IGNORECASE)
            if m_context:
                current_q['content'] += '\n[Context] ' + text
            else:
                current_q['content'] += '\n' + text

    return title, answers, questions


# ─── Build final question list ────────────────────────────────────────────────

def build_db_questions(
    raw_questions: List[dict],
    answers: Dict[str, Dict[str, str]],
    subject_id: int,
) -> List[dict]:
    """Chuyển raw questions sang format DB.

    Returns list of dicts ready for db.insert_question().
    """
    db_qs = []

    for rq in raw_questions:
        section = rq['section']
        num = rq['number']
        content = rq['content'].strip()
        opts = rq.get('opts', {})
        stmts = rq.get('stmts', {})

        # Skip empty
        if not content:
            continue

        if section == 'I':
            question_number = num + TN_OFFSET
            question_type = 'trac_nghiem'
            # Options: from opts buffer, or empty dict
            if opts:
                options = {k: opts[k] for k in 'ABCD' if k in opts}
            else:
                options = {}
            correct_answer = answers['I'].get(str(num), None)

        elif section == 'II':
            question_number = num + DS_OFFSET
            question_type = 'dung_sai'
            # Options: map stmts a/b/c/d → A/B/C/D
            options = {}
            for k, letter in [('a', 'A'), ('b', 'B'), ('c', 'C'), ('d', 'D')]:
                if k in stmts:
                    options[letter] = stmts[k]
                elif opts.get(letter):
                    options[letter] = opts[letter]
            correct_answer = answers['II'].get(str(num), None)

        elif section == 'III':
            question_number = num + TLN_OFFSET
            question_type = 'tu_luan'
            options = None
            correct_answer = answers['III'].get(str(num), None)

        else:
            continue

        # Detect formulas (superscripts, math symbols, ...)
        has_formula = bool(re.search(
            r'[²³⁰¹⁴⁵⁶⁷⁸⁹₀₁₂₃₄₅₆₇₈₉αβγδΔλμπΩ°]|'
            r'(?:phương trình|công thức|tích phân|\bsin\b|\bcos\b|\btan\b)',
            content, re.IGNORECASE
        ))

        db_qs.append({
            'question_number': question_number,
            'question_type': question_type,
            'content': content,
            'options': options if options else None,
            'correct_answer': correct_answer,
            'has_formula': has_formula,
            'subject_id': subject_id,
        })

    return db_qs


# ─── Auto-hide ────────────────────────────────────────────────────────────────

def auto_hide_buggy(exam_id: int, conn) -> int:
    """Ẩn câu hỏi bị lỗi:
    - Nội dung quá ngắn (<15 ký tự)
    - TN không có options (trừ khi chỉ có công thức - chấp nhận)
    - Số câu trùng lặp
    """
    with conn.cursor() as cur:
        hidden = 0

        # 1. Nội dung quá ngắn — nhưng không ẩn TN có đủ 4 đáp án A/B/C/D
        #    (fill-in-blank question như "________" vẫn hợp lệ nếu có options)
        cur.execute("""
            UPDATE questions SET is_hidden = true
            WHERE exam_id = %s AND is_hidden = false AND LENGTH(TRIM(content)) < 15
              AND NOT (
                question_type = 'trac_nghiem'
                AND options IS NOT NULL
                AND (options)::jsonb ? 'A'
                AND (options)::jsonb ? 'B'
                AND (options)::jsonb ? 'C'
                AND (options)::jsonb ? 'D'
              )
        """, (exam_id,))
        hidden += cur.rowcount

        # 2. TN không có options
        cur.execute("""
            UPDATE questions SET is_hidden = true
            WHERE exam_id = %s AND is_hidden = false
              AND question_type = 'trac_nghiem'
              AND (options IS NULL OR options = '{}')
        """, (exam_id,))
        hidden += cur.rowcount

        # 3. DS không có options (statements)
        cur.execute("""
            UPDATE questions SET is_hidden = true
            WHERE exam_id = %s AND is_hidden = false
              AND question_type = 'dung_sai'
              AND (options IS NULL OR options = '{}')
        """, (exam_id,))
        hidden += cur.rowcount

        # 4. Số câu trùng lặp (giữ câu đầu tiên)
        cur.execute("""
            UPDATE questions SET is_hidden = true
            WHERE exam_id = %s AND id NOT IN (
                SELECT MIN(id) FROM questions
                WHERE exam_id = %s
                GROUP BY question_number
            )
        """, (exam_id, exam_id))
        hidden += cur.rowcount

    return hidden


# ─── Main pipeline ────────────────────────────────────────────────────────────

def run_docx_pipeline(
    docx_path: str,
    title: str,
    year: int,
    subject_id: int,
    exam_type: str = 'thi_thu',
) -> int:
    """Parse DOCX và import vào DB.

    Returns:
        exam_id (int)
    """
    console.print(Panel(
        f"[bold]{title}[/]\n[dim]{docx_path}[/]",
        title="[cyan]DOCX Pipeline[/]",
        expand=False,
    ))

    # 1. Parse DOCX
    console.print("  Bước 1: Parse DOCX...")
    raw_title, answers, raw_qs = parse_docx_questions(docx_path)

    console.print(f"    Tiêu đề: [dim]{raw_title[:80]}[/]")
    section_counts = {}
    for q in raw_qs:
        s = q['section']
        section_counts[s] = section_counts.get(s, 0) + 1
    console.print(f"    Câu hỏi raw: {section_counts}")
    console.print(f"    Đáp án: I={len(answers['I'])} II={len(answers['II'])} III={len(answers['III'])}")

    # 2. Build DB questions
    db_qs = build_db_questions(raw_qs, answers, subject_id)
    console.print(f"  Bước 2: {len(db_qs)} câu hợp lệ")

    if len(db_qs) < 10:
        raise ValueError(f"Quá ít câu ({len(db_qs)}), kiểm tra lại DOCX")

    # 3. Tạo exam record
    console.print("  Bước 3: Tạo exam record...")
    with db.get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO exams (title, year, exam_type, subject_id, ocr_status, created_at)
                VALUES (%s, %s, %s, %s, 'done', NOW())
                RETURNING id
            """, (title, year, exam_type, subject_id))
            exam_id = cur.fetchone()[0]
    console.print(f"    exam_id = {exam_id}")

    # 4. Insert questions
    console.print(f"  Bước 4: Insert {len(db_qs)} câu...")
    ok = 0; fail = 0
    for q in db_qs:
        try:
            db.insert_question(
                exam_id=exam_id,
                subject_id=q['subject_id'],
                topic_id=None,
                question_number=q['question_number'],
                content=q['content'],
                content_raw=q['content'],
                question_type=q['question_type'],
                level='Nhận biết',
                level_confidence=0.5,
                options=q['options'],
                correct_answer=q['correct_answer'],
                explanation=None,
                has_formula=q['has_formula'],
                has_image=False,
                has_table=False,
                classification_meta=None,
            )
            ok += 1
        except Exception as e:
            console.print(f"    [red]Lỗi insert câu {q['question_number']}: {e}[/]")
            fail += 1

    console.print(f"    OK: {ok} | Fail: {fail}")

    # 5. Auto-hide buggy
    console.print("  Bước 5: Auto-hide lỗi...")
    with db.get_conn() as conn:
        hidden = auto_hide_buggy(exam_id, conn)
    console.print(f"    Ẩn {hidden} câu lỗi")

    # Final stats
    with db.get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT COUNT(*) FILTER (WHERE NOT is_hidden),
                       COUNT(*) FROM questions WHERE exam_id = %s
            """, (exam_id,))
            vis, total = cur.fetchone()

    console.print(f"  [green]✓ exam_id={exam_id} | {vis} câu visible / {total} total | ẩn {total-vis}[/]")
    return exam_id


# ─── CLI ──────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    import sys
    if len(sys.argv) < 4:
        print("Usage: python import_exam_docx.py <docx_path> <subject_id> <year> [title]")
        print("  subject_id: 2=Lý, 3=Hóa, 6=Sử, 9=Anh")
        sys.exit(1)

    docx_path = sys.argv[1]
    subject_id = int(sys.argv[2])
    year = int(sys.argv[3])
    title = sys.argv[4] if len(sys.argv) > 4 else Path(docx_path).stem

    eid = run_docx_pipeline(docx_path, title, year, subject_id)
    print(f"Done: exam_id={eid}")
