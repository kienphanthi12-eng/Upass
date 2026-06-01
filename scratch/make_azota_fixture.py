"""
scratch/make_azota_fixture.py — Tạo file .docx mẫu ĐÚNG format Azota để test pipeline.

Tạo: data/pdfs/fixture_azota_ly.docx
  - Header "Mã đề 0201"
  - PHẦN I: 3 câu trắc nghiệm (Câu 1 gạch chân đáp án B — LỆCH với bảng để test needs_review)
  - PHẦN II: 1 câu đúng/sai (gạch chân ý a, c → DSDS)
  - PHẦN III: 2 câu trả lời ngắn ("Đáp án: ...")
  - 1 câu có tiền tố mức độ [1, TH]
  - "----- HẾT -----" + BẢNG ĐÁP ÁN (2 mã đề 0201/0202, đáp án P1 khác nhau)
  - Vài "Lời giải"
"""
import sys, io
from pathlib import Path
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document

OUT = Path(__file__).parent.parent / "data" / "pdfs" / "fixture_azota_ly.docx"


def add_options(doc, opts: dict, correct_underline: str | None = None):
    """Mỗi option 1 paragraph; gạch chân nội dung option đúng nếu chỉ định."""
    for letter, content in opts.items():
        p = doc.add_paragraph()
        p.add_run(f"{letter}. ")
        run = p.add_run(content)
        if correct_underline and letter == correct_underline:
            run.underline = True


def add_subitems(doc, items: dict, true_letters: set):
    """Phần II: ý a) b) c) d); gạch chân nội dung ý ĐÚNG."""
    for letter, content in items.items():
        p = doc.add_paragraph()
        p.add_run(f"{letter}) ")
        run = p.add_run(content)
        if letter in true_letters:
            run.underline = True


def main():
    doc = Document()
    doc.add_paragraph("SỞ GD&ĐT LÂM ĐỒNG — ĐỀ THI THỬ TỐT NGHIỆP THPT 2026")
    doc.add_paragraph("Môn thi: VẬT LÍ — Thời gian: 50 phút")
    doc.add_paragraph("Họ và tên: ………………………  Số báo danh: ………  Mã đề 0201")

    # ── PHẦN I ──────────────────────────────────────────────────────────────
    doc.add_paragraph("PHẦN I. Thí sinh trả lời từ câu 1 đến câu 3. Mỗi câu chọn một phương án.")

    doc.add_paragraph("Câu 1. Đơn vị của công suất trong hệ SI là")
    add_options(doc, {"A": "jun (J).", "B": "oát (W).", "C": "niutơn (N).", "D": "paxcan (Pa)."},
                correct_underline="B")  # gạch chân B — nhưng bảng 0201 sẽ ghi A (test mismatch)

    doc.add_paragraph("Câu 2. [1, TH] Nhiệt độ không tuyệt đối ứng với")
    add_options(doc, {"A": "0 K.", "B": "273 K.", "C": "273 °C.", "D": "100 °C."})

    doc.add_paragraph("Câu 3. Đại lượng đặc trưng cho mức quán tính của vật là")
    add_options(doc, {"A": "khối lượng.", "B": "trọng lượng.", "C": "thể tích.", "D": "vận tốc."})
    doc.add_paragraph("Lời giải. Khối lượng là đại lượng đặc trưng cho mức quán tính.")

    # ── PHẦN II ─────────────────────────────────────────────────────────────
    doc.add_paragraph("PHẦN II. Thí sinh trả lời đúng/sai từng ý a), b), c), d).")
    doc.add_paragraph("Câu 1. Cho mạch điện gồm điện trở R và nguồn điện. Xét các phát biểu:")
    add_subitems(doc, {
        "a": "Cường độ dòng điện tỉ lệ thuận với hiệu điện thế.",
        "b": "Điện trở phụ thuộc cường độ dòng điện.",
        "c": "Công suất tỏa nhiệt là P = I²R.",
        "d": "Đơn vị điện trở là culông.",
    }, true_letters={"a", "c"})  # → DSDS

    # ── PHẦN III ────────────────────────────────────────────────────────────
    doc.add_paragraph("PHẦN III. Thí sinh trả lời ngắn.")
    doc.add_paragraph("Câu 1. Một vật khối lượng 2 kg, gia tốc 3 m/s². Lực tác dụng (N) là bao nhiêu?")
    doc.add_paragraph("Đáp án: 6")
    doc.add_paragraph("Câu 2. Tính nhiệt lượng (J) để đun nóng (làm tròn). Đáp án: 4200 | 4200,0")

    # ── HẾT + BẢNG ĐÁP ÁN ───────────────────────────────────────────────────
    doc.add_paragraph("----- HẾT -----")
    doc.add_paragraph("BẢNG ĐÁP ÁN")
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Mã đề"; hdr[1].text = "1"; hdr[2].text = "2"; hdr[3].text = "3"
    r1 = table.rows[1].cells
    r1[0].text = "0201"; r1[1].text = "A"; r1[2].text = "A"; r1[3].text = "A"
    r2 = table.rows[2].cells
    r2[0].text = "0202"; r2[1].text = "D"; r2[2].text = "C"; r2[3].text = "B"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Đã tạo fixture: {OUT}")


if __name__ == "__main__":
    main()
