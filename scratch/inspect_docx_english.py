import docx
import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

def main():
    docx_path = Path("data/pdfs/thuvienhoclieu.com-De-thi-thu-Tot-Nghiep-2026-Tieng-Anh-THPT-Tay-Ninh.docx")
    if not docx_path.exists():
        print(f"File not found: {docx_path}")
        return
        
    doc = docx.Document(docx_path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fullText.append(cell.text)
                
    text = "\n".join(fullText)
    print(f"DOCX text length: {len(text)}")
    
    print("\n=== Searching for 'Question 32' in DOCX ===")
    matches32 = list(re.finditer(r"(Question\s*32)", text, re.I))
    print(f"Found {len(matches32)} matches for Question 32:")
    for m in matches32:
        start = max(0, m.start() - 100)
        end = min(len(text), m.end() + 100)
        print(f"  Match at {m.start()}:\n{text[start:end]!r}\n" + "-"*30)
        
    print("\n=== Searching for 'Question 38' in DOCX ===")
    matches38 = list(re.finditer(r"(Question\s*38)", text, re.I))
    print(f"Found {len(matches38)} matches for Question 38:")
    for m in matches38:
        start = max(0, m.start() - 100)
        end = min(len(text), m.end() + 100)
        print(f"  Match at {m.start()}:\n{text[start:end]!r}\n" + "-"*30)

import re
if __name__ == "__main__":
    main()
